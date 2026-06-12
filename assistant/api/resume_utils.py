import asyncio
import concurrent.futures
from assistant.utils.logger import logger
import pdfplumber
from docx import Document
import os
import json
import io
import base64
import tempfile
from pathlib import Path
from datetime import datetime
from sqlalchemy.orm import Session
from sqlalchemy.exc import IntegrityError
from fastapi import BackgroundTasks
from assistant.enums.user_enum import UserRole
from assistant.LLM.llm_resume_analysis import sync_analyze_resume_with_llm, analyze_resume_with_llm
from assistant.entity import Resume, ResumeStatus, ResumeEducation, ResumeWorkExperience, ResumeSkill, ResumeProject, User

# 尝试导入 PyMuPDF（fitz）用于 PDF 转图片
try:
    import fitz
    HAS_PYMUPDF = True
except ImportError:
    HAS_PYMUPDF = False
    logger.warning("PyMuPDF 未安装，将使用传统文本提取方式")

# 尝试导入 LLM 客户端（用于图片识别）
try:
    from openai import AsyncOpenAI
    from assistant.config.config_manager import ConfigManager
    config_manager = ConfigManager()
    llm_config = config_manager.get_llm_config()
    llm_client = AsyncOpenAI(
        api_key=llm_config.get('api_key'),
        base_url=llm_config.get('url'),
    )
    VISION_MODEL = llm_config.get('model', 'doubao-seed-1-6-251015')
except Exception as e:
    logger.warning(f"LLM 客户端初始化失败: {e}")
    llm_client = None
    VISION_MODEL = None

# 延迟导入 TosFileManager，避免循环依赖
def get_file_manager():
    from assistant.file.file_manager import TosFileManager
    return TosFileManager()


def extract_text_from_pdf(pdf_path):
    """
    使用 pdfplumber 从 PDF 文件中提取文本
    
    Args:
        pdf_path: PDF 文件路径
    
    Returns:
        提取的文本内容
    """
    text = ""
    try:
        with pdfplumber.open(pdf_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
    except Exception as e:
        print(f"解析 PDF 文件失败: {e}")
    return text


def extract_text_from_docx(docx_path):
    """
    从 Word 文件中提取文本
    
    Args:
        docx_path: Word 文件路径
    
    Returns:
        提取的文本内容
    """
    text = ""
    try:
        doc = Document(docx_path)
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
    except Exception as e:
        print(f"解析 Word 文件失败: {e}")
    return text


def extract_text(file_path: str, file_bytes: bytes):
    """
    根据文件类型提取文本
    
    Args:
        file_path: 文件路径
        file_bytes: 文件内容
    
    Returns:
        提取的文本内容
    """
    file_extension = os.path.splitext(file_path)[1].lower()
    try:
        # 2. 按文件类型提取文本
        if file_extension == ".pdf":
            return extract_text_from_pdf_stream(file_bytes)
        elif file_extension in [".docx", ".doc"]:
            return extract_text_from_docx_stream(file_bytes)
        else:
            # 文本文件：直接解码
            return file_bytes.decode("utf-8", errors="replace")
    except Exception as e:
        logger.error(f"简历文本提取失败: {e}")
        return ""

def extract_text_from_pdf_stream(pdf_bytes: bytes) -> str:
    """从PDF二进制流提取文本（PyPDF2实现）"""
    from PyPDF2 import PdfReader
    reader = PdfReader(io.BytesIO(pdf_bytes))
    return "".join([page.extract_text() or "" for page in reader.pages])


def extract_text_from_docx_stream(docx_bytes: bytes) -> str:
    """从Word二进制流提取文本"""
    from docx import Document
    doc = Document(io.BytesIO(docx_bytes))
    return "\n".join([para.text for para in doc.paragraphs])


# ========== PDF 转图片 + 图片识别相关函数 ==========

def pdf_to_images(pdf_bytes: bytes, output_dir: str = None, dpi: int = 200) -> tuple:
    """
    将 PDF 二进制流转为图片
    
    Args:
        pdf_bytes: PDF 文件二进制数据
        output_dir: 图片输出目录
        dpi: 图片分辨率
    
    Returns:
        tuple: (图片路径列表, 输出目录路径)
    """
    if not HAS_PYMUPDF:
        raise ImportError("PyMuPDF 未安装，请先安装: pip install pymupdf")
    
    # 生成临时输出目录
    if output_dir:
        output_dir = Path(output_dir)
    else:
        output_dir = Path(tempfile.gettempdir()) / f"resume_images_{datetime.now().strftime('%Y%m%d_%H%M%S_%f')}"
    
    output_dir.mkdir(parents=True, exist_ok=True)
    
    # 使用 PyMuPDF 打开 PDF
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    num_pages = len(doc)

    def render_page(page_num):
        """渲染单页 PDF 为图片（线程安全，get_pixmap 释放 GIL）"""
        page = doc[page_num]
        pix = page.get_pixmap(matrix=fitz.Matrix(dpi/72, dpi/72))
        image_path = output_dir / f"page_{page_num + 1:03d}.jpg"
        pix.save(str(image_path))
        return str(image_path)

    # 多页 PDF 并发渲染
    if num_pages <= 1:
        image_paths = [render_page(0)]
    else:
        with concurrent.futures.ThreadPoolExecutor(max_workers=min(num_pages, 8)) as executor:
            image_paths = list(executor.map(render_page, range(num_pages)))

    doc.close()
    logger.info(f"PDF 转图片完成，共 {len(image_paths)} 页，输出目录: {output_dir}")
    return image_paths, str(output_dir)


async def extract_text_from_images(image_paths: list) -> str:
    """
    使用视觉 LLM 从图片中提取文本内容（只提取，不做结构化解析）
    
    Args:
        image_paths: 图片路径列表
    
    Returns:
        str: 提取的纯文本内容
    """
    if not llm_client:
        logger.warning("LLM 客户端未初始化，无法使用图片识别")
        return ""
    
    if not image_paths:
        return ""
    
    prompt = """请仔细分析这张简历图片，提取图片中的所有文字内容。
要求：
1. 保留原有的段落结构和换行
2. 准确识别所有中文、英文、数字
3. 不要遗漏任何文字
4. 只返回提取到的文本内容，不要其他解释"""
    
    all_texts = []
    
    for i, image_path in enumerate(image_paths, 1):
        try:
            # 将图片转为 base64
            with open(image_path, "rb") as image_file:
                base64_image = base64.b64encode(image_file.read()).decode("utf-8")
            
            # 构造消息
            messages = [
                {"role": "system", "content": "你是一个专业的简历图片识别助手"},
                {
                    "role": "user",
                    "content": [
                        {"type": "text", "text": prompt},
                        {
                            "type": "image_url",
                            "image_url": {
                                "url": f"data:image/jpeg;base64,{base64_image}"
                            }
                        }
                    ]
                }
            ]
            
            # 调用 LLM 提取文本
            response = await llm_client.chat.completions.create(
                model=VISION_MODEL,
                messages=messages,
                extra_body={"thinking": {"type": "disabled"}},
                max_tokens=4000,
                stream=False,
            )
            
            text = response.choices[0].message.content
            all_texts.append(text)
            logger.info(f"第 {i} 页图片文本提取完成，长度: {len(text)} 字符")
            
        except Exception as e:
            logger.error(f"识别图片失败 {image_path}: {e}")
            all_texts.append(f"[第 {i} 页识别失败: {str(e)}]")
    
    # 合并所有页面的文本
    full_text = "\n\n".join(all_texts)
    logger.info(f"图片文本提取完成，总长度: {len(full_text)} 字符")
    return full_text


async def extract_text_from_images_concurrent(image_paths: list) -> str:
    """
    并发提取所有页面的图片文本（比串行快 N 倍）

    每页独立调 LLM 提取文本，结果按页序合并。
    用于单个页面 token 太大不适用合并模式时的备用方案。
    """
    if not llm_client or not image_paths:
        return ""

    prompt = """请仔细分析这张简历图片，提取图片中的所有文字内容。
要求：
1. 保留原有的段落结构和换行
2. 准确识别所有中文、英文、数字
3. 不要遗漏任何文字
4. 只返回提取到的文本内容，不要其他解释"""

    async def _extract_one(image_path: str, page_num: int) -> str:
        try:
            with open(image_path, "rb") as image_file:
                base64_image = base64.b64encode(image_file.read()).decode("utf-8")

            messages = [
                {"role": "system", "content": "你是一个专业的简历图片识别助手"},
                {
                    "role": "user",
                    "content": [
                        {"type": "text", "text": prompt},
                        {
                            "type": "image_url",
                            "image_url": {"url": f"data:image/jpeg;base64,{base64_image}"}
                        }
                    ]
                }
            ]

            response = await llm_client.chat.completions.create(
                model=VISION_MODEL,
                messages=messages,
                extra_body={"thinking": {"type": "disabled"}},
                max_tokens=4000,
                stream=False,
            )

            text = response.choices[0].message.content
            logger.info(f"第 {page_num} 页并发提取完成，长度: {len(text)} 字符")
            return text
        except Exception as e:
            logger.error(f"识别第 {page_num} 页失败: {e}")
            return f"[第 {page_num} 页识别失败]"

    tasks = [_extract_one(p, i) for i, p in enumerate(image_paths, 1)]
    results = await asyncio.gather(*tasks)

    full_text = "\n\n".join(results)
    logger.info(f"并发提取完成，共 {len(results)} 页，总长度: {len(full_text)} 字符")
    return full_text


async def analyze_resume_with_llm_from_images(image_paths: list) -> dict:
    """
    使用视觉 LLM 直接从简历图片解析出结构化 JSON
    （合并"提取文本" + "结构化解析"两次 LLM 调用为一次）

    将所有图片页一次性发给 LLM，直接输出结构化 JSON，省去中间文本提取步骤。

    Args:
        image_paths: 图片路径列表

    Returns:
        Dict: 结构化解析结果，同 analyze_resume_with_llm 返回格式
    """
    if not llm_client or not image_paths:
        return {"educations": [], "work_experiences": [], "skills": [], "projects": []}

    from assistant.LLM.llm_resume_analysis import extract_json_safe

    prompt = """请仔细分析这些图片，判断其内容是否为简历，如果是则提取所有信息并直接输出 JSON。

强制要求（必须遵守）：
1. 只能输出 JSON，不允许任何额外文字
2. JSON 必须能被 Python 的 json.loads() 正确解析
3. 所有 key 必须使用双引号
4. 所有字符串必须使用双引号
5. 不允许出现多余逗号
6. 不存在的值必须为 null（不能是 "" 或 "无"）
7. 年龄 age 必须是整数，无法确定则为 null
8. is_985 / is_211 必须是整数 0 或 1，无法判断则为 null
9. 日期必须为 "YYYY-MM-DD"，无法确定则为 null

输出 JSON 格式如下：

{
  "is_resume": true,
  "person_info": {
    "name": "string or null",
    "age": number or null,
    "gender": "string or null",
    "phone": "string or null",
    "email": "string or null",
    "address": "string or null"
  },
  "educations": [
    {
      "school_name": "string or null",
      "degree": "string or null",
      "major": "string or null",
      "start_date": "YYYY-MM-DD or null",
      "end_date": "YYYY-MM-DD or null",
      "is_985": 0 or 1 or null,
      "is_211": 0 or 1 or null
    }
  ],
  "work_experiences": [
    {
      "company_name": "string or null",
      "position": "string or null",
      "start_date": "YYYY-MM-DD or null",
      "end_date": "YYYY-MM-DD or null",
      "description": "string or null"
    }
  ],
  "skills": [
    {
      "skill_name": "string or null",
      "proficiency_level": "string or null"
    }
  ],
  "projects": [
    {
      "project_name": "string or null",
      "description": "string or null",
      "start_date": "YYYY-MM-DD or null",
      "end_date": "YYYY-MM-DD or null",
      "role": "string or null"
    }
  ]
}

注意：最终输出必须是一个纯 JSON 字符串，不允许有任何前后缀内容或 markdown 标记。"""

    # 构造消息——把所有图片放在同一次请求中
    content = [{"type": "text", "text": prompt}]
    for image_path in image_paths:
        with open(image_path, "rb") as f:
            base64_image = base64.b64encode(f.read()).decode("utf-8")
        content.append({
            "type": "image_url",
            "image_url": {"url": f"data:image/jpeg;base64,{base64_image}"}
        })

    messages = [
        {"role": "system", "content": "你是一个专业的简历解析助手"},
        {"role": "user", "content": content}
    ]

    try:
        response = await llm_client.chat.completions.create(
            model=VISION_MODEL,
            messages=messages,
            extra_body={"thinking": {"type": "disabled"}},
            max_tokens=4000,
            stream=False,
        )

        response_content = response.choices[0].message.content
        parsed_data = extract_json_safe(response_content)
        logger.info(f"简历图片直接解析结果: {parsed_data}")
        return parsed_data
    except Exception as e:
        logger.error(f"简历图片直接解析失败: {e}")
        return {"educations": [], "work_experiences": [], "skills": [], "projects": []}


async def analyze_resume_only(file_bytes: bytes, filename: str) -> dict:
    """
    纯 LLM 分析，不涉及数据库操作。
    返回分析结果 dict（包含 is_resume、person_info、educations 等字段），失败返回 {}。
    """
    file_extension = os.path.splitext(filename)[1].lower()
    image_paths = None
    output_dir = None

    try:
        if file_extension == ".pdf" and HAS_PYMUPDF:
            logger.info(f"文件 {filename} 为 PDF，使用图片识别模式")
            loop = asyncio.get_event_loop()
            image_paths, output_dir = await loop.run_in_executor(
                None, lambda: pdf_to_images(file_bytes)
            )
            parsed_data = await analyze_resume_with_llm_from_images(image_paths)
            return parsed_data if parsed_data else {}
        else:
            logger.info(f"文件 {filename} 为 {file_extension}，使用传统文本提取模式")
            loop = asyncio.get_event_loop()
            resume_text = await loop.run_in_executor(None, lambda: extract_text(filename, file_bytes))
            parsed_data = await analyze_resume_with_llm(resume_text)
            return parsed_data if parsed_data else {}
    finally:
        if image_paths:
            try:
                cleanup_temp_images(image_paths, output_dir)
            except Exception:
                pass


def cleanup_temp_images(image_paths: list, output_dir: str = None):
    """
    清理临时图片文件和目录（文件不存在则静默跳过）
    
    Args:
        image_paths: 图片路径列表
        output_dir: 输出目录（可选，删除图片后尝试删除空目录）
    """
    # 先删除图片文件
    for img_path in image_paths:
        try:
            if os.path.exists(img_path):
                os.remove(img_path)
                logger.debug(f"已删除临时图片: {img_path}")
        except Exception as e:
            logger.warning(f"删除临时图片失败 {img_path}: {e}")
    
    # 尝试删除空目录
    if output_dir and os.path.exists(output_dir):
        try:
            os.rmdir(output_dir)
            logger.debug(f"已删除临时目录: {output_dir}")
        except Exception as e:
            logger.warning(f"删除临时目录失败 {output_dir}: {e}（目录可能非空）")
    
    logger.info("临时图片清理完成")


def delete_resume_file(file_path: str, db: Session = None):
    """后台删除简历文件"""
    try:
        file_manager = get_file_manager()
        file_manager.delete_file(file_path, db)
    except Exception as e:
        logger.error(f"删除文件失败: {e}")


def delete_resume_data(
    resume_id: int,
    db: Session,
    current_user_id: int = 0,
    skip_background: bool = False
):
    """
    删除简历数据（数据库层面，不含 API 路由）
    
    Args:
        resume_id: 简历 ID
        db: 数据库会话
        current_user_id: 当前用户 ID（0 表示跳过权限检查）
        skip_background: 是否跳过后台文件删除
    """
    # 根据简历ID查找简历
    db_resume = db.query(Resume).filter(Resume.id == resume_id).first()
    if not db_resume:
        raise ValueError("简历不存在")
    
    # 验证权限（如果提供了 current_user_id）
    if current_user_id and current_user_id > 0:
        recruiter = db.query(User).filter(User.id == current_user_id).first()
        if not recruiter:
            raise ValueError("用户不存在")
        if db_resume.user_id != current_user_id:
            raise ValueError("简历不存在或不属于当前用户所有")
    
    # 保存文件路径用于后台删除
    file_path = db_resume.file_path
    
    # 删除关联表数据
    db.query(ResumeEducation).filter(ResumeEducation.resume_id == db_resume.id).delete()
    db.query(ResumeWorkExperience).filter(ResumeWorkExperience.resume_id == db_resume.id).delete()
    db.query(ResumeSkill).filter(ResumeSkill.resume_id == db_resume.id).delete()
    db.query(ResumeProject).filter(ResumeProject.resume_id == db_resume.id).delete()
    if db_resume.candidate_name and db_resume.candidate_name != "待解析":
        db.query(User).filter(User.username == db_resume.candidate_name).delete()
    
    # 删除主表记录
    db.delete(db_resume)
    db.commit()
    
    # 后台删除文件
    if not skip_background:
        delete_resume_file(file_path, db)
    
    return True


def parse_resume(file_path):
    """
    解析简历文件
    
    Args:
        file_path: 简历文件路径
    
    Returns:
        解析后的结构化数据
    """
    # 提取文本
    text = extract_text(file_path)
    if not text:
        print("提取文本失败")
        return None
    
    print("提取的文本内容:", text[:500], "..." if len(text) > 500 else "")
    
    # 使用 LLM 解析文本
    parsed_data = sync_analyze_resume_with_llm(text)
    if not parsed_data:
        print("LLM 解析失败")
        return None
    
    print("解析结果 (JSON 格式):", json.dumps(parsed_data, ensure_ascii=False, indent=2))
    
    return parsed_data


def store_resume_details(db, resume_id, parsed_data, current_user_id: int):
    """
    存储简历详情到相关表
    
    Args:
        db: 数据库会话
        resume_id: 简历 ID
        parsed_data: 解析后的结构化数据
    """
    person_info = parsed_data.get("person_info", {})

    # 容错：LLM 有时返回的 JSON 缺少 person_info 外层，字段直接挂在顶层
    if not person_info and parsed_data.get("name"):
        person_info = {
            "name": parsed_data.get("name"),
            "age": parsed_data.get("age"),
            "gender": parsed_data.get("gender"),
            "phone": parsed_data.get("phone"),
            "email": parsed_data.get("email"),
            "address": parsed_data.get("address"),
        }

    username = None

    if person_info:
        # Validate and truncate fields to ensure they fit in database columns
        username = (person_info.get("name") or "")[:50]
        email = (person_info.get("email") or "")[:100]
        # Remove non-digit characters and truncate to 15 digits (reasonable phone number length)
        phone = ''.join(filter(str.isdigit, person_info.get("phone") or ""))[:15]

        # 同时检查 username 和 email，只要有一个匹配就算已存在
        db_user = db.query(User).filter(
            (User.username == username) | (User.email == email)
        ).first()
        
        if db_user:
            # 用户已存在，更新信息
            db_user.email = email
            db_user.recruiter_id = current_user_id
            db_user.phone = phone
            db.commit()
            logger.info(f"[用户已存在，更新成功] username: {username}, email: {email}")
        else:
            # 新用户，创建
            db_user = User(
                username=username,
                recruiter_id=current_user_id,
                email=email,
                phone=phone,
                role=UserRole.CANDIDATE,
                status="CREATED",
            )
            db.add(db_user)
            try:
                db.commit()  # 立即提交，避免后续外键依赖问题
                db.refresh(db_user)
                logger.info(f"[新用户创建成功] username: {username}, email: {email}")
            except IntegrityError:
                # 并发场景：另一个后台任务已创建同邮箱/同名的用户
                db.rollback()
                db_user = db.query(User).filter(
                    (User.username == username) | (User.email == email)
                ).first()
                if db_user:
                    db_user.email = email
                    db_user.recruiter_id = current_user_id
                    db_user.phone = phone
                    db.commit()
                    db.refresh(db_user)
                    logger.info(f"[用户已存在，更新成功(并发)] username: {username}, email: {email}")
    
    # 先删除该简历原有的关联记录（防止重复导入导致重复数据）
    db.query(ResumeEducation).filter(ResumeEducation.resume_id == resume_id).delete()
    db.query(ResumeWorkExperience).filter(ResumeWorkExperience.resume_id == resume_id).delete()
    db.query(ResumeSkill).filter(ResumeSkill.resume_id == resume_id).delete()
    db.query(ResumeProject).filter(ResumeProject.resume_id == resume_id).delete()

    # 存储教育经历
    for edu in parsed_data.get("educations", []):
        # 处理日期字段，确保不为 None
        start_date_str = edu.get("start_date", "2023-01-01")
        end_date_str = edu.get("end_date", "2023-12-31")
        
        try:
            start_date = datetime.strptime(start_date_str, "%Y-%m-%d") if start_date_str else datetime.now()
            end_date = datetime.strptime(end_date_str, "%Y-%m-%d") if end_date_str else datetime.now()
        except (ValueError, TypeError):
            start_date = datetime.now()
            end_date = datetime.now()
        
        db_education = ResumeEducation(
            resume_id=resume_id,
            school_name=(edu.get("school_name") or "")[:100],
            degree=(edu.get("degree") or "")[:50],
            major=(edu.get("major") or "")[:100],
            start_date=start_date,
            end_date=end_date,
            is_985=edu.get("is_985", 0),
            is_211=edu.get("is_211", 0)
        )
        db.add(db_education)
    
    # 存储工作经历
    for work in parsed_data.get("work_experiences", []):
        # 处理日期字段，确保不为 None
        start_date_str = work.get("start_date", "2023-01-01")
        end_date_str = work.get("end_date", "2023-12-31")
        
        try:
            start_date = datetime.strptime(start_date_str, "%Y-%m-%d") if start_date_str else datetime.now()
            end_date = datetime.strptime(end_date_str, "%Y-%m-%d") if end_date_str else datetime.now()
        except (ValueError, TypeError):
            start_date = datetime.now()
            end_date = datetime.now()
        
        db_work = ResumeWorkExperience(
            resume_id=resume_id,
            company_name=(work.get("company_name") or "")[:100],
            position=(work.get("position") or "")[:100],
            start_date=start_date,
            end_date=end_date,
            description=work.get("description") or ""
        )
        db.add(db_work)
    
    # 存储技能
    for skill in parsed_data.get("skills", []):
        db_skill = ResumeSkill(
            resume_id=resume_id,
            skill_name=(skill.get("skill_name") or "")[:100],
            proficiency_level=(skill.get("proficiency_level") or "")[:20]
        )
        db.add(db_skill)
    
    # 存储项目经历
    for project in parsed_data.get("projects", []):
        # 处理日期字段，确保不为 None
        start_date_str = project.get("start_date", "2023-01-01")
        end_date_str = project.get("end_date", "2023-12-31")
        
        try:
            start_date = datetime.strptime(start_date_str, "%Y-%m-%d") if start_date_str else datetime.now()
            end_date = datetime.strptime(end_date_str, "%Y-%m-%d") if end_date_str else datetime.now()
        except (ValueError, TypeError):
            start_date = datetime.now()
            end_date = datetime.now()
        
        db_project = ResumeProject(
            resume_id=resume_id,
            project_name=(project.get("project_name") or "")[:100],
            description=project.get("description") or "",
            start_date=start_date,
            end_date=end_date,
            role=(project.get("role") or "")[:100]
        )
        db.add(db_project)
    
    db.commit()
    if username:
        return username
    else:
        return "未知"


def save_resume_to_db(
    db: Session,
    resume_id: int,
    parsed_data: dict,
    current_user_id: int,
    tos_key: str = None,
    filename: str = "",
) -> bool:
    """
    将 LLM 分析结果写入数据库（纯同步函数，在 asyncio.to_thread 中执行）。
    """
    resume = db.query(Resume).filter(Resume.id == resume_id).first()
    if not resume:
        logger.error(f"简历{resume_id}不存在")
        return False

    # 非简历 → 清理数据
    if parsed_data and not parsed_data.get("is_resume", True):
        logger.warning(f"LLM 判定非简历，清理数据: resume_id={resume_id}, file={filename}")
        db.query(ResumeEducation).filter(ResumeEducation.resume_id == resume_id).delete()
        db.query(ResumeWorkExperience).filter(ResumeWorkExperience.resume_id == resume_id).delete()
        db.query(ResumeSkill).filter(ResumeSkill.resume_id == resume_id).delete()
        db.query(ResumeProject).filter(ResumeProject.resume_id == resume_id).delete()
        db.query(Resume).filter(Resume.id == resume_id).delete()
        if tos_key:
            try:
                get_file_manager().delete_file(tos_key, db)
            except Exception as e:
                logger.warning(f"删除 TOS 文件失败: {e}")
        db.commit()
        logger.info(f"非简历文件已清理: resume_id={resume_id}")
        return False

    # 有效简历：存储解析结果
    if parsed_data:
        candidate_name = store_resume_details(db, resume_id, parsed_data, current_user_id)
        resume = db.query(Resume).filter(Resume.id == resume_id).first()

        # 查重：若已存在同名简历，合并到现有记录
        if candidate_name and candidate_name != "待解析":
            existing = db.query(Resume).filter(
                Resume.user_id == current_user_id,
                Resume.candidate_name == candidate_name,
                Resume.id != resume_id
            ).first()

            if existing:
                logger.info(f"检测到同名简历: {candidate_name} (现有ID={existing.id}, 新ID={resume_id})，执行覆盖")
                old_tos_key = existing.file_path

                existing.file_path = resume.file_path
                existing.file_type = resume.file_type
                existing.content = resume.content
                existing.status = ResumeStatus.ANALYZED
                existing.extracted_at = datetime.now()
                existing.original_file_name = resume.original_file_name

                db.query(ResumeEducation).filter(ResumeEducation.resume_id == existing.id).delete()
                db.query(ResumeWorkExperience).filter(ResumeWorkExperience.resume_id == existing.id).delete()
                db.query(ResumeSkill).filter(ResumeSkill.resume_id == existing.id).delete()
                db.query(ResumeProject).filter(ResumeProject.resume_id == existing.id).delete()

                for edu in db.query(ResumeEducation).filter(ResumeEducation.resume_id == resume_id).all():
                    edu.resume_id = existing.id
                for work in db.query(ResumeWorkExperience).filter(ResumeWorkExperience.resume_id == resume_id).all():
                    work.resume_id = existing.id
                for skill in db.query(ResumeSkill).filter(ResumeSkill.resume_id == resume_id).all():
                    skill.resume_id = existing.id
                for proj in db.query(ResumeProject).filter(ResumeProject.resume_id == resume_id).all():
                    proj.resume_id = existing.id

                db.delete(resume)
                db.commit()

                if old_tos_key:
                    try:
                        get_file_manager().delete_file(old_tos_key, db)
                    except Exception as e:
                        logger.warning(f"删除旧 TOS 文件失败: {e}")

                logger.info(f"简历 {candidate_name} 已合并到 ID={existing.id}，新 ID={resume_id} 已删除")
                return True

        # 无重复，正常更新当前简历
        resume = db.query(Resume).filter(Resume.id == resume_id).first()
        resume.status = ResumeStatus.ANALYZED
        resume.extracted_at = datetime.now()
        resume.candidate_name = candidate_name[:50] if candidate_name else "待解析"
        db.commit()
        logger.info(f"简历 {resume_id} 分析完成")
        return True
    else:
        resume.status = ResumeStatus.FAILED_ANALYSIS
        db.commit()
        logger.error(f"简历 {resume_id} 分析结果为空")
        return False


def process_resume(db, user_id, file_path, file_type):
    """
    处理简历文件，包括解析和存储
    
    Args:
        db: 数据库会话
        user_id: 用户 ID
        file_path: 简历文件路径
        file_type: 文件类型
    
    Returns:
        简历 ID 和解析结果
    """
    try:
        # 提取文本
        content = extract_text(file_path)
        
        # 解析简历
        parsed_data = parse_resume(file_path)
        if not parsed_data:
            print("解析失败")
            return None, None
        
        return None, parsed_data
    except Exception as e:
        print(f"处理简历失败: {e}")
        db.rollback()
        return None, None


async def process_resume_background(db, resume_id, resume_text: str, current_user_id: int, image_paths: list = None):
    """
    后台处理简历分析（支持传统文本提取和图片识别两种方式）
    
    Args:
        db: 数据库会话
        resume_id: 简历 ID
        resume_text: 传统方式提取的文本（备用）
        current_user_id: 当前用户 ID
        image_paths: 图片路径列表（可选，有则优先使用图片识别）
    """
    try:
        # 1. 优先使用图片识别，否则使用传统文本提取
        final_text = resume_text
        use_image_recognition = False
        
        if image_paths and len(image_paths) > 0:
            logger.info(f"使用图片识别模式处理简历 {resume_id}，共 {len(image_paths)} 页")
            try:
                final_text = await extract_text_from_images(image_paths)
                use_image_recognition = True
            except Exception as e:
                logger.warning(f"图片识别失败，回退到传统文本提取: {e}")
        
        # 2. 调用 analyze_resume_with_llm 进行结构化解析
        parsed_data = await analyze_resume_with_llm(final_text)

        # 3. 查询简历记录
        resume = db.query(Resume).filter(Resume.id == resume_id).first()
        if not resume:
            logger.error(f"后台任务错误：简历{resume_id}不存在")
            return
        
        # 4. 存储解析结果
        if parsed_data:
            candidate_name = store_resume_details(db, resume_id, parsed_data, current_user_id)
            resume.status = ResumeStatus.ANALYZED
            resume.extracted_at = datetime.now()
            resume.candidate_name = candidate_name[:50] if candidate_name else "待解析"
            resume.content = final_text  # 存储识别到的文本用于追溯
            db.commit()
            logger.info(f"简历 {resume_id} 分析完成（{'图片识别' if use_image_recognition else '文本提取'}）")
        else:
            resume.status = ResumeStatus.FAILED_ANALYSIS
            db.commit()
            logger.error(f"简历 {resume_id} 分析失败")
    
    except Exception as e:
        logger.error(f"简历{resume_id}后台分析失败: {str(e)}", exc_info=True)  # 打印完整堆栈
        resume = db.query(Resume).filter(Resume.id == resume_id).first()
        if resume:
            resume.status = ResumeStatus.FAILED_ANALYSIS
            db.commit()
        db.rollback()


async def process_resume_background_with_images(db, resume_id, file_bytes: bytes, filename: str, current_user_id: int, tos_key: str = None) -> bool:
    """
    后台处理简历分析（PDF 转图片后识别）

    保留此函数对外兼容（import_resume 端点使用），
    内部委托给 analyze_resume_only + save_resume_to_db。
    """
    try:
        parsed_data = await analyze_resume_only(file_bytes, filename)
        return save_resume_to_db(db, resume_id, parsed_data, current_user_id, tos_key, filename)
    except Exception as e:
        logger.error(f"简历{resume_id}后台分析失败: {str(e)}", exc_info=True)
        resume = db.query(Resume).filter(Resume.id == resume_id).first()
        if resume:
            resume.status = ResumeStatus.FAILED_ANALYSIS
            db.commit()
        db.rollback()
        return False

